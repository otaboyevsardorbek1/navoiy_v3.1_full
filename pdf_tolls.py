#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import os
import sys
import re
import json
import csv
import hashlib
import base64
import threading
import time
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor, as_completed
from datetime import datetime
from collections import Counter
from typing import Dict, List, Tuple, Optional, Any, Union
from dataclasses import dataclass, field
from enum import Enum

# Asosiy PDF kutubxonalari
try:
    import pdfplumber
except ImportError:
    print("❌ pdfplumber o'rnatilmagan. O'rnatish: pip install pdfplumber")
    sys.exit(1)

# Qo'shimcha kutubxonalar (optional)
try:
    from cryptography.fernet import Fernet
    CRYPTO_AVAILABLE = True
except ImportError:
    CRYPTO_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from openpyxl import Workbook
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from googletrans import Translator
    TRANSLATOR_AVAILABLE = True
except ImportError:
    TRANSLATOR_AVAILABLE = False

try:
    from pdf2image import convert_from_path
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# Pydantic uchun (Navoiy API formati)
try:
    from pydantic import BaseModel
    PYDANTIC_AVAILABLE = True
except ImportError:
    PYDANTIC_AVAILABLE = False
    # Pydantic bo'lmasa, oddiy klass yaratamiz
    class BaseModel:
        def __init__(self, **kwargs):
            for k, v in kwargs.items():
                setattr(self, k, v)
        
        def model_dump_json(self, **kwargs):
            import json
            data = {k: v for k, v in self.__dict__.items() if not k.startswith('_')}
            return json.dumps(data, ensure_ascii=False, **kwargs)


# ==================== RANGLI CHIQISH UCHUN ====================
class Colors:
    """Terminal ranglari"""
    HEADER = '\033[95m'
    BLUE = '\033[94m'
    CYAN = '\033[96m'
    GREEN = '\033[92m'
    YELLOW = '\033[93m'
    RED = '\033[91m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'
    END = '\033[0m'
    GRAY = '\033[90m'


def print_progress_bar(iteration, total, prefix='', suffix='', length=50, fill='█'):
    """Progress barni chiqarish"""
    percent = f"{100 * (iteration / float(total)):.1f}"
    filled_length = int(length * iteration // total)
    bar = fill * filled_length + '░' * (length - filled_length)
    
    sys.stdout.write(f'\r{prefix} |{Colors.CYAN}{bar}{Colors.END}| {percent}% {suffix}')
    sys.stdout.flush()
    
    if iteration == total:
        print()


class Spinner:
    """Yuklanish animatsiyasi"""
    def __init__(self, message="Ishlanmoqda"):
        self.message = message
        self.spinning = False
        self.spinner_chars = ['⠋', '⠙', '⠹', '⠸', '⠼', '⠴', '⠦', '⠧', '⠇', '⠏']
        self.idx = 0
        
    def spin(self):
        while self.spinning:
            sys.stdout.write(f'\r{Colors.CYAN}{self.spinner_chars[self.idx]}{Colors.END} {self.message}')
            sys.stdout.flush()
            self.idx = (self.idx + 1) % len(self.spinner_chars)
            time.sleep(0.1)
    
    def start(self):
        self.spinning = True
        threading.Thread(target=self.spin, daemon=True).start()
    
    def stop(self, success_message="✓ Bajarildi"):
        self.spinning = False
        time.sleep(0.2)
        sys.stdout.write(f'\r{Colors.GREEN}✓{Colors.END} {success_message}                    \n')
        sys.stdout.flush()


# ==================== DATA KLASSLAR ====================
@dataclass
class PageData:
    """Sahifa ma'lumotlari"""
    page_num: int
    text: str
    char_count: int = 0
    word_count: int = 0
    tables: List[Any] = field(default_factory=list)
    images_count: int = 0
    
    def __post_init__(self):
        self.char_count = len(self.text)
        self.word_count = len(self.text.split())


@dataclass
class ExtractionResult:
    """Ekstraksiya natijasi"""
    file_path: str
    file_name: str
    file_size_kb: float
    total_pages: int
    pages: List[PageData] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    extraction_time: float = 0.0
    extracted_at: str = field(default_factory=lambda: datetime.now().isoformat())
    
    @property
    def full_text(self) -> str:
        return "\n\n".join([f"{'='*50}\nSAHIFA {p.page_num}\n{'='*50}\n\n{p.text}" for p in self.pages])
    
    @property
    def total_chars(self) -> int:
        return sum(p.char_count for p in self.pages)
    
    @property
    def total_words(self) -> int:
        return sum(p.word_count for p in self.pages)


# ==================== TEXT PROCESSOR ====================
class TextProcessor:
    """Matnni qayta ishlash klassi"""
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Matnni tozalash"""
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s\.\,\!\?\-\:\;\(\)\[\]\{\}\@\#\$\%\&\*\+\=\/\\\'\"\№\«\»\—\–]', '', text)
        return text.strip()
    
    @staticmethod
    def find_emails(text: str) -> List[str]:
        pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        return list(set(re.findall(pattern, text)))
    
    @staticmethod
    def find_phone_numbers(text: str) -> List[str]:
        patterns = [
            r'\+998[\s\-]?\d{2}[\s\-]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}',
            r'\b\d{2}[\s\-]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}\b',
            r'\+998[\s\-]?\(\d{2}\)[\s\-]?\d{3}[\-\s]?\d{2}[\-\s]?\d{2}'
        ]
        phones = []
        for pattern in patterns:
            phones.extend(re.findall(pattern, text))
        return list(set(phones))
    
    @staticmethod
    def find_urls(text: str) -> List[str]:
        pattern = r'https?://[^\s<>"{}|\\^`\[\]]+|www\.[^\s<>"{}|\\^`\[\]]+'
        return list(set(re.findall(pattern, text)))
    
    @staticmethod
    def find_dates(text: str) -> List[str]:
        patterns = [
            r'\d{2}\.\d{2}\.\d{4}',
            r'\d{4}-\d{2}-\d{2}',
            r'\d{2}/\d{2}/\d{4}',
            r'\d{1,2}\s+(yanvar|fevral|mart|aprel|may|iyun|iyul|avgust|sentabr|oktabr|noyabr|dekabr)\s+\d{4}'
        ]
        dates = []
        for pattern in patterns:
            dates.extend(re.findall(pattern, text, re.IGNORECASE))
        return dates
    
    @staticmethod
    def word_frequency(text: str, top_n: int = 20) -> List[Tuple[str, int]]:
        words = re.findall(r'\b[a-zA-Zа-яА-ЯўЎқҚғҒҳҲ]{3,}\b', text.lower())
        return Counter(words).most_common(top_n)
    
    @staticmethod
    def summarize(text: str, max_sentences: int = 5) -> str:
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
        
        if len(sentences) <= max_sentences:
            return '. '.join(sentences) + '.'
        
        important = sorted(sentences, key=lambda x: len(x.split()), reverse=True)[:max_sentences]
        return '. '.join(important) + '.'
    
    @staticmethod
    def translate(text: str, target_lang: str = 'uz') -> Optional[str]:
        if not TRANSLATOR_AVAILABLE:
            return None
        
        try:
            translator = Translator()
            result = translator.translate(text[:5000], dest=target_lang)
            return result.text
        except Exception:
            return None
    
    @staticmethod
    def extract_keywords(text: str, top_n: int = 10) -> List[str]:
        stop_words = {'va', 'bilan', 'uchun', 'ham', 'deb', 'edi', 'ekan', 'bo\'lib', 
                      'the', 'and', 'for', 'with', 'this', 'that', 'from', 'have'}
        
        words = re.findall(r'\b[a-zA-Zа-яА-ЯўЎқҚғҒҳҲ]{4,}\b', text.lower())
        filtered = [w for w in words if w not in stop_words]
        return [word for word, _ in Counter(filtered).most_common(top_n)]


# ==================== SEARCH ENGINE ====================
class SearchEngine:
    """Qidiruv tizimi"""
    
    def __init__(self, result: ExtractionResult):
        self.result = result
        self._index = None
    
    def _build_index(self):
        self._index = {}
        for page in self.result.pages:
            words = set(re.findall(r'\b\w+\b', page.text.lower()))
            for word in words:
                if word not in self._index:
                    self._index[word] = []
                self._index[word].append(page.page_num)
    
    def search(self, query: str, case_sensitive: bool = False) -> List[Dict]:
        results = []
        flags = 0 if case_sensitive else re.IGNORECASE
        
        for page in self.result.pages:
            if re.search(query, page.text, flags):
                matches = list(re.finditer(query, page.text, flags))
                for match in matches:
                    start = max(0, match.start() - 50)
                    end = min(len(page.text), match.end() + 50)
                    context = page.text[start:end]
                    
                    results.append({
                        'page': page.page_num,
                        'position': match.start(),
                        'context': f"...{context}...",
                        'match': match.group()
                    })
        
        return results
    
    def search_regex(self, pattern: str) -> List[Dict]:
        results = []
        try:
            for page in self.result.pages:
                matches = re.finditer(pattern, page.text)
                for match in matches:
                    results.append({
                        'page': page.page_num,
                        'match': match.group(),
                        'groups': match.groups() if match.groups() else None
                    })
        except re.error as e:
            return [{'error': str(e)}]
        
        return results
    
    def get_statistics(self) -> Dict:
        return {
            'total_pages': self.result.total_pages,
            'total_chars': self.result.total_chars,
            'total_words': self.result.total_words,
            'avg_chars_per_page': self.result.total_chars / max(1, self.result.total_pages),
            'avg_words_per_page': self.result.total_words / max(1, self.result.total_pages),
            'file_size_kb': self.result.file_size_kb
        }


# ==================== PDF EXTRACTOR (PARALLEL) ====================
class ParallelPDFExtractor:
    """Parallel PDF ekstraktor"""
    
    def __init__(self, pdf_path: str, max_workers: int = 4):
        self.pdf_path = pdf_path
        self.max_workers = max_workers
        self.result = ExtractionResult(
            file_path=pdf_path,
            file_name=os.path.basename(pdf_path),
            file_size_kb=os.path.getsize(pdf_path) / 1024 if os.path.exists(pdf_path) else 0,
            total_pages=0
        )
    
    def get_info(self) -> str:
        try:
            with pdfplumber.open(self.pdf_path) as pdf:
                self.result.total_pages = len(pdf.pages)
                self.result.metadata = pdf.metadata if pdf.metadata else {}
                
                info = f"""
{Colors.BOLD}{Colors.CYAN}╔══════════════════════════════════════════════════════════════╗
║                         PDF MA'LUMOTLARI                        ║
╚══════════════════════════════════════════════════════════════╝{Colors.END}
{Colors.GREEN}📁 Fayl nomi:{Colors.END} {self.result.file_name}
{Colors.GREEN}📁 To'liq yo'l:{Colors.END} {self.pdf_path}
{Colors.GREEN}📏 Hajmi:{Colors.END} {self.result.file_size_kb:.2f} KB ({self.result.file_size_kb/1024:.2f} MB)
{Colors.GREEN}📄 Sahifalar soni:{Colors.END} {self.result.total_pages}
"""
                if self.result.metadata:
                    info += f"\n{Colors.YELLOW}📋 Metadata:{Colors.END}\n"
                    for key, value in self.result.metadata.items():
                        if value and str(value).strip():
                            info += f"   {Colors.GRAY}{key}:{Colors.END} {value}\n"
                
                return info
        except Exception as e:
            return f"{Colors.RED}❌ Xatolik: {str(e)}{Colors.END}"
    
    def _extract_single_page(self, page_num: int, pdf_path: str) -> Optional[PageData]:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                if page_num <= len(pdf.pages):
                    page = pdf.pages[page_num - 1]
                    text = page.extract_text() or ""
                    tables = page.extract_tables() or []
                    
                    return PageData(
                        page_num=page_num,
                        text=TextProcessor.clean_text(text),
                        tables=tables
                    )
        except Exception:
            return None
        return None
    
    def extract_all_pages(self, show_progress: bool = True) -> ExtractionResult:
        start_time = time.time()
        
        spinner = Spinner("PDF tahlil qilinmoqda...")
        if show_progress:
            spinner.start()
        
        try:
            with pdfplumber.open(self.pdf_path) as pdf:
                self.result.total_pages = len(pdf.pages)
                self.result.metadata = pdf.metadata or {}
        except Exception as e:
            if show_progress:
                spinner.stop(f"❌ Xatolik: {e}")
            return self.result
        
        if show_progress:
            spinner.stop(f"Jami {self.result.total_pages} sahifa topildi")
        
        print(f"\n{Colors.CYAN}⏳ Sahifalar parallel qayta ishlanmoqda...{Colors.END}")
        
        pages_data = []
        completed = 0
        
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            futures = {
                executor.submit(self._extract_single_page, i, self.pdf_path): i 
                for i in range(1, self.result.total_pages + 1)
            }
            
            for future in as_completed(futures):
                page_num = futures[future]
                try:
                    page_data = future.result(timeout=30)
                    if page_data:
                        pages_data.append(page_data)
                except Exception:
                    pass
                
                completed += 1
                if show_progress:
                    print_progress_bar(
                        completed, self.result.total_pages,
                        prefix=f'{Colors.BOLD}Progress:{Colors.END}',
                        suffix=f'{completed}/{self.result.total_pages} sahifa'
                    )
        
        self.result.pages = sorted(pages_data, key=lambda x: x.page_num)
        self.result.extraction_time = time.time() - start_time
        
        print(f"\n{Colors.GREEN}✅ Ekstraksiya yakunlandi!{Colors.END}")
        print(f"{Colors.GRAY}⏱️ Vaqt: {self.result.extraction_time:.2f} soniya{Colors.END}")
        print(f"{Colors.GRAY}📊 O'rtacha tezlik: {self.result.total_pages/self.result.extraction_time:.1f} sahifa/sekund{Colors.END}")
        
        return self.result
    
    def extract_range(self, start_page: int, end_page: int) -> ExtractionResult:
        start_time = time.time()
        
        pages_data = []
        total_to_extract = end_page - start_page + 1
        
        print(f"\n{Colors.CYAN}⏳ {start_page}-{end_page} sahifalar olinmoqda...{Colors.END}")
        
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            futures = {
                executor.submit(self._extract_single_page, i, self.pdf_path): i 
                for i in range(start_page, end_page + 1)
            }
            
            completed = 0
            for future in as_completed(futures):
                try:
                    page_data = future.result(timeout=30)
                    if page_data:
                        pages_data.append(page_data)
                except Exception:
                    pass
                
                completed += 1
                print_progress_bar(
                    completed, total_to_extract,
                    prefix=f'{Colors.BOLD}Progress:{Colors.END}',
                    suffix=f'{completed}/{total_to_extract}'
                )
        
        self.result.pages = sorted(pages_data, key=lambda x: x.page_num)
        self.result.extraction_time = time.time() - start_time
        
        return self.result


# ==================== OCR EXTRACTOR ====================
class OCRExtractor:
    """Rasmli PDF uchun OCR"""
    
    def __init__(self):
        self.available = OCR_AVAILABLE
        if not self.available:
            print(f"{Colors.YELLOW}⚠️ OCR funksiyasi uchun kutubxonalar o'rnatilmagan{Colors.END}")
            print(f"{Colors.GRAY}   O'rnatish: pip install pytesseract pdf2image pillow{Colors.END}")
    
    def extract(self, pdf_path: str, languages: str = 'uzb+rus+eng') -> Optional[str]:
        if not self.available:
            return None
        
        spinner = Spinner("PDF rasmlarga aylantirilmoqda...")
        spinner.start()
        
        try:
            images = convert_from_path(pdf_path, dpi=200)
            spinner.stop(f"{len(images)} ta rasmga aylantirildi")
        except Exception as e:
            spinner.stop(f"❌ Xatolik: {e}")
            return None
        
        all_text = ""
        for i, image in enumerate(images, 1):
            print_progress_bar(i, len(images), prefix='OCR:', suffix=f'{i}/{len(images)}')
            try:
                text = pytesseract.image_to_string(image, lang=languages, config='--psm 1')
                all_text += f"\n\n{'='*50}\nSAHIFA {i} (OCR)\n{'='*50}\n\n{text}"
            except Exception as e:
                all_text += f"\n\nSAHIFA {i}: OCR xatosi - {e}\n\n"
        
        return all_text


# ==================== EXPORT FORMATS ====================
class ExportManager:
    """Eksport boshqaruvchisi"""
    
    @staticmethod
    def to_txt(result: ExtractionResult, output_path: str) -> bool:
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(f"# PDF MATNI\n")
                f.write(f"# Fayl: {result.file_name}\n")
                f.write(f"# Olingan sana: {result.extracted_at}\n")
                f.write(f"# Jami sahifalar: {result.total_pages}\n")
                f.write(f"# {'='*50}\n\n")
                f.write(result.full_text)
            return True
        except Exception:
            return False
    
    @staticmethod
    def to_json(result: ExtractionResult, output_path: str) -> bool:
        try:
            data = {
                'metadata': {
                    'file_name': result.file_name,
                    'file_path': result.file_path,
                    'file_size_kb': result.file_size_kb,
                    'total_pages': result.total_pages,
                    'pdf_metadata': result.metadata,
                    'extraction_time': result.extraction_time,
                    'extracted_at': result.extracted_at,
                    'total_chars': result.total_chars,
                    'total_words': result.total_words
                },
                'pages': [
                    {
                        'page_num': p.page_num,
                        'text': p.text,
                        'char_count': p.char_count,
                        'word_count': p.word_count,
                        'tables_count': len(p.tables)
                    }
                    for p in result.pages
                ]
            }
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            return True
        except Exception:
            return False
    
    @staticmethod
    def to_csv(result: ExtractionResult, output_path: str) -> bool:
        try:
            with open(output_path, 'w', encoding='utf-8', newline='') as f:
                writer = csv.writer(f)
                writer.writerow(['Sahifa', 'Belgilar', 'So\'zlar', 'Matn'])
                for p in result.pages:
                    writer.writerow([p.page_num, p.char_count, p.word_count, p.text[:500]])
            return True
        except Exception:
            return False
    
    @staticmethod
    def to_word(result: ExtractionResult, output_path: str) -> bool:
        if not DOCX_AVAILABLE:
            return False
        
        try:
            doc = docx.Document()
            doc.add_heading(f'PDF Matni: {result.file_name}', 0)
            doc.add_paragraph(f'Olingan sana: {result.extracted_at}')
            doc.add_paragraph(f'Jami sahifalar: {result.total_pages}')
            doc.add_page_break()
            
            for p in result.pages:
                doc.add_heading(f'SAHIFA {p.page_num}', level=1)
                doc.add_paragraph(p.text)
                if p.page_num < result.total_pages:
                    doc.add_page_break()
            
            doc.save(output_path)
            return True
        except Exception:
            return False
    
    @staticmethod
    def to_excel(result: ExtractionResult, output_path: str) -> bool:
        if not EXCEL_AVAILABLE:
            return False
        
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "PDF Matni"
            
            ws['A1'] = "Sahifa"
            ws['B1'] = "Belgilar soni"
            ws['C1'] = "So'zlar soni"
            ws['D1'] = "Matn"
            
            for i, p in enumerate(result.pages, 2):
                ws[f'A{i}'] = p.page_num
                ws[f'B{i}'] = p.char_count
                ws[f'C{i}'] = p.word_count
                ws[f'D{i}'] = p.text[:1000]
            
            wb.save(output_path)
            return True
        except Exception:
            return False
    
    @staticmethod
    def to_html(result: ExtractionResult, output_path: str) -> bool:
        try:
            import html
            
            html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>PDF Matni: {result.file_name}</title>
    <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 40px; background: #f5f5f5; }}
        .container {{ max-width: 1200px; margin: 0 auto; background: white; padding: 30px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
        h1 {{ color: #333; border-bottom: 3px solid #4CAF50; padding-bottom: 10px; }}
        .page {{ margin-bottom: 40px; padding: 20px; background: #fafafa; border-radius: 5px; }}
        h2 {{ color: #4CAF50; margin-top: 0; }}
        pre {{ white-space: pre-wrap; word-wrap: break-word; font-family: inherit; }}
        .meta {{ color: #666; font-size: 14px; margin-bottom: 20px; }}
    </style>
</head>
<body>
    <div class="container">
        <h1>📚 PDF Matni: {result.file_name}</h1>
        <div class="meta">
            <p>📄 Sahifalar: {result.total_pages} | 📏 Hajm: {result.file_size_kb:.2f} KB | 📅 Olingan: {result.extracted_at}</p>
        </div>
"""
            for p in result.pages:
                html_content += f"""
        <div class="page">
            <h2>📃 SAHIFA {p.page_num}</h2>
            <pre>{html.escape(p.text)}</pre>
        </div>
"""
            html_content += """
    </div>
</body>
</html>"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            return True
        except Exception:
            return False


# ==================== SECURITY ====================
class SecurityManager:
    """Xavfsizlik boshqaruvchisi"""
    
    @staticmethod
    def generate_key() -> Optional[bytes]:
        if CRYPTO_AVAILABLE:
            return Fernet.generate_key()
        return None
    
    @staticmethod
    def encrypt_file(input_path: str, output_path: str, key: bytes) -> bool:
        if not CRYPTO_AVAILABLE:
            return False
        
        try:
            f = Fernet(key)
            with open(input_path, 'rb') as file:
                file_data = file.read()
            encrypted = f.encrypt(file_data)
            with open(output_path, 'wb') as file:
                file.write(encrypted)
            return True
        except Exception:
            return False
    
    @staticmethod
    def decrypt_file(input_path: str, output_path: str, key: bytes) -> bool:
        if not CRYPTO_AVAILABLE:
            return False
        
        try:
            f = Fernet(key)
            with open(input_path, 'rb') as file:
                encrypted_data = file.read()
            decrypted = f.decrypt(encrypted_data)
            with open(output_path, 'wb') as file:
                file.write(decrypted)
            return True
        except Exception:
            return False
    
    @staticmethod
    def get_file_hash(file_path: str, algorithm: str = 'sha256') -> str:
        hash_func = hashlib.new(algorithm)
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b''):
                hash_func.update(chunk)
        return hash_func.hexdigest()


# ==================== NAVOIY API JSON FORMATI ====================
class QuizTypeEnum(str, Enum):
    """Quiz turlari"""
    single = "single"
    multiple = "multiple"
    text = "text"


class QuizInContent:
    """Sahifa ichidagi quiz (viktorina)"""
    def __init__(self, id: int, question: str, type: QuizTypeEnum = QuizTypeEnum.single,
                 options: List[str] = None, correct_answers: List[int] = None,
                 difficulty: int = 1, points: int = 10, explanation: str = None):
        self.id = id
        self.question = question
        self.type = type
        self.options = options or []
        self.correct_answers = correct_answers or []
        self.difficulty = difficulty
        self.points = points
        self.explanation = explanation
    
    def to_dict(self) -> dict:
        return {
            "id": self.id,
            "question": self.question,
            "type": self.type.value if hasattr(self.type, 'value') else str(self.type),
            "options": self.options,
            "correct_answers": self.correct_answers,
            "difficulty": self.difficulty,
            "points": self.points,
            "explanation": self.explanation
        }


class PageContent:
    """Sahifa kontenti"""
    def __init__(self, page_number: int, title: str = "", word_count: int = 0,
                 content: str = "", quizzes: List[QuizInContent] = None):
        self.page_number = page_number
        self.title = title
        self.word_count = word_count
        self.content = content
        self.quizzes = quizzes or []
    
    def to_dict(self) -> dict:
        return {
            "page_number": self.page_number,
            "title": self.title,
            "word_count": self.word_count,
            "content": self.content,
            "quizzes": [q.to_dict() for q in self.quizzes]
        }


class AsarContentJSON:
    """Asar to'liq kontenti - Navoiy API formati"""
    def __init__(self, asar_id: int = None, slug: str = "", title: str = "",
                 version: int = 1, checksum: str = "", total_pages: int = 0,
                 language: str = "uz_cyrillic", pages: List[PageContent] = None,
                 metadata: dict = None):
        self.asar_id = asar_id
        self.slug = slug
        self.title = title
        self.version = version
        self.checksum = checksum
        self.total_pages = total_pages
        self.language = language
        self.pages = pages or []
        self.metadata = metadata or {}
    
    def to_dict(self, exclude_checksum: bool = False) -> dict:
        data = {
            "asar_id": self.asar_id,
            "slug": self.slug,
            "title": self.title,
            "version": self.version,
            "total_pages": self.total_pages,
            "language": self.language,
            "pages": [p.to_dict() for p in self.pages],
            "metadata": self.metadata
        }
        if not exclude_checksum:
            data["checksum"] = self.checksum
        return data
    
    def model_dump_json(self, indent: int = 2, ensure_ascii: bool = False, exclude: set = None) -> str:
        exclude = exclude or set()
        data = self.to_dict(exclude_checksum=('checksum' in exclude))
        return json.dumps(data, ensure_ascii=ensure_ascii, indent=indent)


class NavoiyJSONExporter:
    """Navoiy API formatida JSON eksport qilish"""
    
    @staticmethod
    def convert_pdf_result_to_navoiy_format(
        result: ExtractionResult,
        asar_slug: str = None,
        asar_title: str = None,
        language: str = "uz_cyrillic",
        auto_generate_quizzes: bool = True
    ) -> AsarContentJSON:
        """PDF natijasini Navoiy API JSON formatiga o'girish"""
        import hashlib
        import re
        
        # Slug va title aniqlash
        if not asar_slug:
            asar_slug = re.sub(r'[^\w\-]', '-', result.file_name.replace('.pdf', '').lower())
        if not asar_title:
            asar_title = result.file_name.replace('.pdf', '').replace('_', ' ').title()
        
        # Sahifalarni PageContent formatiga o'girish
        pages = []
        for page in result.pages:
            # Har bir sahifa uchun quiz generatsiya qilish
            quizzes = []
            if auto_generate_quizzes and len(page.text) > 200:
                quizzes = NavoiyJSONExporter._generate_quizzes_from_text(
                    page.text, 
                    page.page_num,
                    result.total_pages
                )
            
            # Sahifa sarlavhasini aniqlash
            lines = page.text.strip().split('\n')
            title = lines[0][:50] if lines else f"Sahifa {page.page_num}"
            
            pages.append(PageContent(
                page_number=page.page_num,
                title=title,
                word_count=page.word_count,
                content=page.text,
                quizzes=quizzes
            ))
        
        # JSON obyektini yaratish
        content_obj = AsarContentJSON(
            asar_id=None,
            slug=asar_slug,
            title=asar_title,
            version=1,
            checksum="",
            total_pages=len(pages),
            language=language,
            pages=pages,
            metadata={
                "source_file": result.file_name,
                "file_size_kb": result.file_size_kb,
                "extraction_time": result.extraction_time,
                "extracted_at": result.extracted_at,
                "total_chars": result.total_chars,
                "total_words": result.total_words,
                "pdf_metadata": result.metadata
            }
        )
        
        # Checksum hisoblash
        json_str = content_obj.model_dump_json(indent=2, ensure_ascii=False, exclude={'checksum'})
        content_obj.checksum = hashlib.sha256(json_str.encode()).hexdigest()
        
        return content_obj
    
    @staticmethod
    def _generate_quizzes_from_text(text: str, page_num: int, total_pages: int) -> List[QuizInContent]:
        """Matndan avtomatik quiz generatsiya qilish"""
        import re
        
        quizzes = []
        sentences = re.split(r'[.!?۔]+', text)
        sentences = [s.strip() for s in sentences if len(s.strip()) > 30]
        
        if not sentences:
            return quizzes
        
        # 1-quiz: Sahifa soni haqida
        if total_pages > 1:
            quizzes.append(QuizInContent(
                id=len(quizzes) + 1,
                question="Asar jami necha sahifadan iborat?",
                type=QuizTypeEnum.single,
                options=[
                    str(total_pages - 10),
                    str(total_pages),
                    str(total_pages + 5),
                    str(total_pages + 10)
                ],
                correct_answers=[1],
                difficulty=1,
                points=10,
                explanation=f"Asar {total_pages} sahifadan iborat."
            ))
        
        # 2-quiz: Raqamli ma'lumotlardan quiz yaratish
        numbers = re.findall(r'\b(\d+)\s*(yil|yosh|bob|sahifa|qism|son|kishi|ming|yuz)\b', text.lower())
        if numbers:
            for num, unit in numbers[:2]:
                if int(num) > 10:
                    question = f"Matnda keltirilgan {unit} soni qancha?"
                    correct_num = int(num)
                    
                    quizzes.append(QuizInContent(
                        id=len(quizzes) + 1,
                        question=question,
                        type=QuizTypeEnum.single,
                        options=[
                            str(correct_num - 5),
                            str(correct_num),
                            str(correct_num + 3),
                            str(correct_num + 7)
                        ],
                        correct_answers=[1],
                        difficulty=2,
                        points=15,
                        explanation=f"To'g'ri javob: {correct_num}"
                    ))
        
        # 3-quiz: Ismlardan quiz yaratish
        names = re.findall(r'\b([A-Z][a-z]+|[А-Я][а-я]+|[A-Z]{2,})\b', text)
        unique_names = list(set(names))[:4]
        if len(unique_names) >= 1:
            main_name = unique_names[0]
            
            options = unique_names[:4] if len(unique_names) >= 4 else unique_names + ["Aniqlanmadi"] * (4 - len(unique_names))
            
            quizzes.append(QuizInContent(
                id=len(quizzes) + 1,
                question="Matnda tilga olingan asosiy shaxs/qahramon kim?",
                type=QuizTypeEnum.single,
                options=options,
                correct_answers=[0],
                difficulty=2,
                points=15,
                explanation=f"Asosiy qahramon: {main_name}"
            ))
        
        return quizzes[:5]
    
    @staticmethod
    def save_to_json(content_obj: AsarContentJSON, output_path: str, indent: int = 2) -> bool:
        """AsarContentJSON obyektini JSON faylga saqlash"""
        try:
            json_str = content_obj.model_dump_json(indent=indent, ensure_ascii=False, exclude={'checksum'})
            # Checksumni qayta qo'shish
            data = json.loads(json_str)
            data['checksum'] = content_obj.checksum
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=indent)
            return True
        except Exception as e:
            print(f"JSON saqlashda xatolik: {e}")
            return False
    
    @staticmethod
    def save_as_separate_pages(
        result: ExtractionResult,
        output_dir: str,
        asar_slug: str = None,
        asar_title: str = None
    ) -> List[str]:
        """Har bir sahifani alohida JSON fayl sifatida saqlash"""
        saved_files = []
        
        os.makedirs(output_dir, exist_ok=True)
        
        if not asar_slug:
            asar_slug = result.file_name.replace('.pdf', '').replace(' ', '_').lower()
        if not asar_title:
            asar_title = result.file_name.replace('.pdf', '').replace('_', ' ').title()
        
        # Metadata fayl
        metadata = {
            "asar_id": None,
            "slug": asar_slug,
            "title": asar_title,
            "total_pages": result.total_pages,
            "language": "uz_cyrillic",
            "version": 1,
            "source_file": result.file_name,
            "file_size_kb": result.file_size_kb,
            "extraction_time": result.extraction_time,
            "extracted_at": result.extracted_at,
            "total_chars": result.total_chars,
            "total_words": result.total_words
        }
        
        metadata_path = os.path.join(output_dir, f"{asar_slug}_metadata.json")
        with open(metadata_path, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)
        saved_files.append(metadata_path)
        
        # Har bir sahifa uchun alohida JSON
        for page in result.pages:
            quizzes = NavoiyJSONExporter._generate_quizzes_from_text(
                page.text, page.page_num, result.total_pages
            )
            
            page_data = {
                "page_number": page.page_num,
                "title": f"Sahifa {page.page_num}",
                "word_count": page.word_count,
                "content": page.text,
                "quizzes": [q.to_dict() for q in quizzes]
            }
            
            page_path = os.path.join(output_dir, f"{asar_slug}_page_{page.page_num:03d}.json")
            with open(page_path, 'w', encoding='utf-8') as f:
                json.dump(page_data, f, ensure_ascii=False, indent=2)
            saved_files.append(page_path)
        
        # To'liq asar JSON
        full_content = NavoiyJSONExporter.convert_pdf_result_to_navoiy_format(
            result, asar_slug, asar_title
        )
        full_path = os.path.join(output_dir, f"{asar_slug}_full.json")
        NavoiyJSONExporter.save_to_json(full_content, full_path)
        saved_files.append(full_path)
        
        return saved_files


# ==================== ASOSIY MENYU ====================
class PDFExtractorCLI:
    """Asosiy CLI dastur"""
    
    def __init__(self):
        self.current_result: Optional[ExtractionResult] = None
        self.current_pdf_path: Optional[str] = None
        self.max_workers = min(8, os.cpu_count() or 4)
    
    def print_header(self):
        """Sarlavhani chiqarish"""
        os.system('cls' if os.name == 'nt' else 'clear')
        print(f"""
{Colors.CYAN}{Colors.BOLD}
╔══════════════════════════════════════════════════════════════════╗
║                                                                  ║
║         📚 ADVANCED PDF MATN OLISH DASTURI v2.1                 ║
║              Parallel • OCR • Export • Search • Navoiy           ║
║                                                                  ║
╚══════════════════════════════════════════════════════════════════╝
{Colors.END}
{Colors.GRAY}Python 3.8.10 • Max Workers: {self.max_workers} • pdfplumber{Colors.END}
""")
    
    def print_menu(self):
        """Menyuni chiqarish"""
        print(f"""
{Colors.BOLD}📋 ASOSIY MENYU:{Colors.END}
{'-'*50}
{Colors.GREEN} 1.{Colors.END} 📁 PDF fayl tanlash va ma'lumot olish
{Colors.GREEN} 2.{Colors.END} 📄 Barcha sahifalardan matn olish {Colors.CYAN}[Parallel]{Colors.END}
{Colors.GREEN} 3.{Colors.END} 📑 Belgilangan sahifalardan matn olish
{Colors.GREEN} 4.{Colors.END} 🔍 Matn ichida qidiruv
{Colors.GREEN} 5.{Colors.END} 📊 Matn tahlili va statistika
{Colors.GREEN} 6.{Colors.END} 📧 Email/Telefon/URL ajratib olish
{Colors.GREEN} 7.{Colors.END} 💾 Turli formatlarda eksport qilish
{'-'*50}
{Colors.YELLOW} 8.{Colors.END} 🖼️ OCR bilan matn olish {Colors.GRAY}[rasmli PDF]{Colors.END}
{Colors.YELLOW} 9.{Colors.END} 🔐 Shifrlash/Xavfsizlik
{Colors.YELLOW}10.{Colors.END} ⚙️ Sozlamalar (Workerlar soni: {self.max_workers})
{'-'*50}
{Colors.RED} 0.{Colors.END} 🚪 Chiqish
""")
        
        if self.current_pdf_path:
            print(f"{Colors.GREEN}📁 Joriy PDF:{Colors.END} {os.path.basename(self.current_pdf_path)}")
            if self.current_result:
                print(f"{Colors.GREEN}   Sahifalar:{Colors.END} {self.current_result.total_pages} | {Colors.GREEN}Matn hajmi:{Colors.END} {self.current_result.total_chars:,} belgi")
    
    def run(self):
        """Dasturni ishga tushirish"""
        while True:
            self.print_header()
            self.print_menu()
            
            choice = input(f"\n{Colors.BOLD}👉 Tanlang (0-10):{Colors.END} ").strip()
            
            if choice == '0':
                print(f"\n{Colors.GREEN}✅ Dastur yakunlandi. Xayr!{Colors.END}")
                break
            elif choice == '1':
                self.select_pdf()
            elif choice == '2':
                self.extract_all_pages()
            elif choice == '3':
                self.extract_range()
            elif choice == '4':
                self.search_text()
            elif choice == '5':
                self.show_statistics()
            elif choice == '6':
                self.extract_contacts()
            elif choice == '7':
                self.export_menu()
            elif choice == '8':
                self.ocr_extract()
            elif choice == '9':
                self.security_menu()
            elif choice == '10':
                self.settings_menu()
            else:
                print(f"{Colors.RED}❌ Noto'g'ri tanlov!{Colors.END}")
            
            if choice != '0':
                input(f"\n{Colors.GRAY}Davom etish uchun ENTER bosing...{Colors.END}")
    
    def select_pdf(self):
        pdf_path = input(f"\n{Colors.CYAN}📁 PDF fayl manzilini kiriting:{Colors.END} ").strip()
        
        if not pdf_path:
            return
        
        pdf_path = pdf_path.strip('"\'').strip()
        
        if not os.path.exists(pdf_path):
            print(f"{Colors.RED}❌ Fayl topilmadi!{Colors.END}")
            return
        
        self.current_pdf_path = pdf_path
        extractor = ParallelPDFExtractor(pdf_path, self.max_workers)
        print(extractor.get_info())
        self.current_result = None
    
    def extract_all_pages(self):
        if not self.current_pdf_path:
            print(f"{Colors.YELLOW}⚠️ Avval PDF fayl tanlang!{Colors.END}")
            return
        
        extractor = ParallelPDFExtractor(self.current_pdf_path, self.max_workers)
        self.current_result = extractor.extract_all_pages(show_progress=True)
        
        if self.current_result and self.current_result.pages:
            self._save_after_extraction()
    
    def extract_range(self):
        if not self.current_pdf_path:
            print(f"{Colors.YELLOW}⚠️ Avval PDF fayl tanlang!{Colors.END}")
            return
        
        try:
            start = int(input(f"{Colors.CYAN}Boshlang'ich sahifa:{Colors.END} "))
            end = int(input(f"{Colors.CYAN}Oxirgi sahifa:{Colors.END} "))
            
            if start < 1 or end < start:
                print(f"{Colors.RED}❌ Noto'g'ri sahifa raqamlari!{Colors.END}")
                return
            
            extractor = ParallelPDFExtractor(self.current_pdf_path, self.max_workers)
            self.current_result = extractor.extract_range(start, end)
            
            if self.current_result and self.current_result.pages:
                self._save_after_extraction()
                
        except ValueError:
            print(f"{Colors.RED}❌ Noto'g'ri raqam kiritildi!{Colors.END}")
    
    def _save_after_extraction(self):
        if not self.current_result:
            return
        
        print(f"\n{Colors.GREEN}📊 Natijalar:{Colors.END}")
        print(f"   Sahifalar: {len(self.current_result.pages)}")
        print(f"   Jami belgilar: {self.current_result.total_chars:,}")
        print(f"   Jami so'zlar: {self.current_result.total_words:,}")
        
        save = input(f"\n{Colors.CYAN}💾 Matnni faylga saqlash kerakmi? (ha/yo'q):{Colors.END} ").lower()
        
        if save in ['ha', 'h', 'yes', 'y']:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            default_name = f"{os.path.splitext(os.path.basename(self.current_pdf_path))[0]}_extracted_{timestamp}.txt"
            
            output_path = input(f"{Colors.CYAN}Fayl nomi [{default_name}]:{Colors.END} ").strip()
            if not output_path:
                output_path = default_name
            
            if ExportManager.to_txt(self.current_result, output_path):
                print(f"{Colors.GREEN}✅ Saqlandi: {output_path}{Colors.END}")
            else:
                print(f"{Colors.RED}❌ Saqlashda xatolik!{Colors.END}")
    
    def search_text(self):
        if not self.current_result:
            print(f"{Colors.YELLOW}⚠️ Avval matn oling!{Colors.END}")
            return
        
        search = SearchEngine(self.current_result)
        stats = search.get_statistics()
        
        print(f"\n{Colors.BOLD}📊 Statistika:{Colors.END}")
        print(f"   Sahifalar: {stats['total_pages']}")
        print(f"   So'zlar: {stats['total_words']:,}")
        print(f"   Belgilar: {stats['total_chars']:,}")
        
        query = input(f"\n{Colors.CYAN}🔍 Qidiruv so'zi yoki regex pattern:{Colors.END} ").strip()
        
        if not query:
            return
        
        regex_mode = input(f"{Colors.CYAN}Regex rejimi? (ha/yo'q):{Colors.END} ").lower() in ['ha', 'h', 'yes', 'y']
        
        if regex_mode:
            results = search.search_regex(query)
        else:
            case_sensitive = input(f"{Colors.CYAN}Katta-kichik harf sezgir? (ha/yo'q):{Colors.END} ").lower() in ['ha', 'h', 'yes', 'y']
            results = search.search(query, case_sensitive)
        
        if results:
            print(f"\n{Colors.GREEN}✅ {len(results)} ta natija topildi:{Colors.END}\n")
            for i, r in enumerate(results[:20], 1):
                print(f"{Colors.YELLOW}{i}. Sahifa {r.get('page', 'N/A')}:{Colors.END}")
                if 'context' in r:
                    print(f"   {r['context'][:100]}")
                elif 'match' in r:
                    print(f"   Mos: {r['match']}")
                print()
            
            if len(results) > 20:
                print(f"{Colors.GRAY}... yana {len(results)-20} ta natija{Colors.END}")
        else:
            print(f"{Colors.YELLOW}ℹ️ Natija topilmadi.{Colors.END}")
    
    def show_statistics(self):
        if not self.current_result:
            print(f"{Colors.YELLOW}⚠️ Avval matn oling!{Colors.END}")
            return
        
        search = SearchEngine(self.current_result)
        stats = search.get_statistics()
        full_text = self.current_result.full_text
        
        print(f"\n{Colors.BOLD}{Colors.CYAN}╔══════════════════════════════════════════════════════════════╗")
        print(f"║                      MATN STATISTIKASI                        ║")
        print(f"╚══════════════════════════════════════════════════════════════╝{Colors.END}\n")
        
        print(f"{Colors.GREEN}📊 Umumiy ma'lumotlar:{Colors.END}")
        print(f"   Sahifalar soni: {stats['total_pages']}")
        print(f"   Jami belgilar: {stats['total_chars']:,}")
        print(f"   Jami so'zlar: {stats['total_words']:,}")
        print(f"   O'rtacha belgi/sahifa: {stats['avg_chars_per_page']:.0f}")
        print(f"   O'rtacha so'z/sahifa: {stats['avg_words_per_page']:.0f}")
        
        print(f"\n{Colors.GREEN}📈 So'z chastotasi (TOP 20):{Colors.END}")
        freq = TextProcessor.word_frequency(full_text, 20)
        for i, (word, count) in enumerate(freq, 1):
            bar_length = int(count / max(1, freq[0][1]) * 30)
            bar = '█' * bar_length
            print(f"   {i:2}. {word:20} {Colors.CYAN}{bar}{Colors.END} {count}")
        
        print(f"\n{Colors.GREEN}🏷️ Kalit so'zlar:{Colors.END}")
        keywords = TextProcessor.extract_keywords(full_text, 15)
        print(f"   {', '.join(keywords)}")
    
    def extract_contacts(self):
        if not self.current_result:
            print(f"{Colors.YELLOW}⚠️ Avval matn oling!{Colors.END}")
            return
        
        full_text = self.current_result.full_text
        
        print(f"\n{Colors.BOLD}{Colors.CYAN}📧 KONTAKT MA'LUMOTLAR{Colors.END}\n")
        
        emails = TextProcessor.find_emails(full_text)
        if emails:
            print(f"{Colors.GREEN}📧 Email manzillar ({len(emails)}):{Colors.END}")
            for email in emails:
                print(f"   • {email}")
        
        phones = TextProcessor.find_phone_numbers(full_text)
        if phones:
            print(f"\n{Colors.GREEN}📞 Telefon raqamlar ({len(phones)}):{Colors.END}")
            for phone in phones[:20]:
                print(f"   • {phone}")
        
        urls = TextProcessor.find_urls(full_text)
        if urls:
            print(f"\n{Colors.GREEN}🌐 URL manzillar ({len(urls)}):{Colors.END}")
            for url in urls[:20]:
                print(f"   • {url}")
        
        dates = TextProcessor.find_dates(full_text)
        if dates:
            print(f"\n{Colors.GREEN}📅 Sanalar ({len(dates)}):{Colors.END}")
            for date in dates[:20]:
                print(f"   • {date}")
        
        if not any([emails, phones, urls, dates]):
            print(f"{Colors.YELLOW}ℹ️ Kontakt ma'lumotlar topilmadi.{Colors.END}")
    
    def export_navoiy_format(self):
        """Navoiy API JSON formatida eksport qilish"""
        if not self.current_result:
            print(f"{Colors.YELLOW}⚠️ Avval matn oling!{Colors.END}")
            return
        
        print(f"\n{Colors.BOLD}{Colors.CYAN}╔══════════════════════════════════════════════════════════════╗")
        print(f"║              NAVOIY API JSON FORMATIDA EKSPORT                ║")
        print(f"╚══════════════════════════════════════════════════════════════╝{Colors.END}\n")
        
        # Asar ma'lumotlarini kiritish
        default_slug = re.sub(r'[^\w\-]', '-', self.current_result.file_name.replace('.pdf', '').lower())
        default_title = self.current_result.file_name.replace('.pdf', '').replace('_', ' ').replace('-', ' ').title()
        
        asar_slug = input(f"{Colors.CYAN}Asar slug [{default_slug}]:{Colors.END} ").strip()
        if not asar_slug:
            asar_slug = default_slug
        
        asar_title = input(f"{Colors.CYAN}Asar nomi [{default_title}]:{Colors.END} ").strip()
        if not asar_title:
            asar_title = default_title
        
        language = input(f"{Colors.CYAN}Til [uz_cyrillic]:{Colors.END} ").strip()
        if not language:
            language = "uz_cyrillic"
        
        print(f"\n{Colors.BOLD}Eksport turi:{Colors.END}")
        print(f"{Colors.GREEN}1.{Colors.END} Yagona JSON fayl (barcha sahifalar bitta faylda)")
        print(f"{Colors.GREEN}2.{Colors.END} Alohida sahifalar (har bir sahifa alohida JSON)")
        print(f"{Colors.GREEN}3.{Colors.END} Ikkala format ham")
        
        export_type = input(f"\n{Colors.CYAN}👉 Tanlang (1-3):{Colors.END} ").strip()
        
        # Quiz generatsiya so'rash
        auto_quiz = input(f"{Colors.CYAN}Avtomatik quiz generatsiya qilinsinmi? (ha/yo'q) [ha]:{Colors.END} ").lower()
        auto_quiz = auto_quiz not in ['yo\'q', 'y', 'no', 'n']
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        try:
            if export_type == '1':
                # Yagona JSON fayl
                content_obj = NavoiyJSONExporter.convert_pdf_result_to_navoiy_format(
                    self.current_result,
                    asar_slug=asar_slug,
                    asar_title=asar_title,
                    language=language,
                    auto_generate_quizzes=auto_quiz
                )
                
                output_path = f"{asar_slug}_{timestamp}.json"
                
                spinner = Spinner("JSON formatida saqlanmoqda...")
                spinner.start()
                
                if NavoiyJSONExporter.save_to_json(content_obj, output_path):
                    spinner.stop(f"✅ Saqlandi: {output_path}")
                    
                    # Statistika
                    print(f"\n{Colors.GREEN}📊 Eksport statistikasi:{Colors.END}")
                    print(f"   Sahifalar: {content_obj.total_pages}")
                    print(f"   Jami quizlar: {sum(len(p.quizzes) for p in content_obj.pages)}")
                    print(f"   Checksum: {content_obj.checksum[:16]}...")
                else:
                    spinner.stop(f"❌ Saqlashda xatolik!")
            
            elif export_type == '2':
                # Alohida sahifalar
                output_dir = f"{asar_slug}_pages_{timestamp}"
                
                spinner = Spinner(f"Sahifalar {output_dir}/ katalogiga saqlanmoqda...")
                spinner.start()
                
                saved_files = NavoiyJSONExporter.save_as_separate_pages(
                    self.current_result,
                    output_dir,
                    asar_slug=asar_slug,
                    asar_title=asar_title
                )
                
                spinner.stop(f"✅ {len(saved_files)} ta fayl saqlandi: {output_dir}/")
                
                print(f"\n{Colors.GREEN}📁 Saqlangan fayllar:{Colors.END}")
                for f in saved_files[:5]:
                    print(f"   • {os.path.basename(f)}")
                if len(saved_files) > 5:
                    print(f"   ... yana {len(saved_files)-5} ta fayl")
            
            elif export_type == '3':
                # Ikkala format ham
                output_dir = f"{asar_slug}_export_{timestamp}"
                os.makedirs(output_dir, exist_ok=True)
                
                spinner = Spinner("Ikkala formatda saqlanmoqda...")
                spinner.start()
                
                # Yagona JSON
                content_obj = NavoiyJSONExporter.convert_pdf_result_to_navoiy_format(
                    self.current_result,
                    asar_slug=asar_slug,
                    asar_title=asar_title,
                    language=language,
                    auto_generate_quizzes=auto_quiz
                )
                full_path = os.path.join(output_dir, f"{asar_slug}_full.json")
                NavoiyJSONExporter.save_to_json(content_obj, full_path)
                
                # Alohida sahifalar
                pages_dir = os.path.join(output_dir, "pages")
                saved_files = NavoiyJSONExporter.save_as_separate_pages(
                    self.current_result,
                    pages_dir,
                    asar_slug=asar_slug,
                    asar_title=asar_title
                )
                
                spinner.stop(f"✅ Hammasi saqlandi: {output_dir}/")
                
                print(f"\n{Colors.GREEN}📁 Saqlangan:{Colors.END}")
                print(f"   • To'liq JSON: {asar_slug}_full.json")
                print(f"   • Sahifalar: pages/ ({len(saved_files)-2} ta)")
                print(f"   • Metadata: {asar_slug}_metadata.json")
            
            else:
                print(f"{Colors.RED}❌ Noto'g'ri tanlov!{Colors.END}")
                
        except Exception as e:
            print(f"{Colors.RED}❌ Eksportda xatolik: {e}{Colors.END}")
            import traceback
            traceback.print_exc()
    
    def export_menu(self):
        """Kengaytirilgan eksport menyusi"""
        if not self.current_result:
            print(f"{Colors.YELLOW}⚠️ Avval matn oling!{Colors.END}")
            return
        
        print(f"\n{Colors.BOLD}💾 EKSPORT FORMATLARI:{Colors.END}")
        print(f"{Colors.GREEN}1.{Colors.END} TXT - Oddiy matn")
        print(f"{Colors.GREEN}2.{Colors.END} JSON - Strukturalangan")
        print(f"{Colors.GREEN}3.{Colors.END} CSV - Jadval ko'rinishida")
        print(f"{Colors.GREEN}4.{Colors.END} HTML - Veb sahifa")
        
        if DOCX_AVAILABLE:
            print(f"{Colors.GREEN}5.{Colors.END} DOCX - Microsoft Word")
        else:
            print(f"{Colors.GRAY}5. DOCX - [o'rnatilmagan]{Colors.END}")
        
        if EXCEL_AVAILABLE:
            print(f"{Colors.GREEN}6.{Colors.END} XLSX - Microsoft Excel")
        else:
            print(f"{Colors.GRAY}6. XLSX - [o'rnatilmagan]{Colors.END}")
        
        print(f"{Colors.CYAN}{Colors.BOLD}7.{Colors.END} 📚 NAVOIY API JSON FORMATI {Colors.GREEN}[YANGI]{Colors.END}")
        
        choice = input(f"\n{Colors.CYAN}👉 Format tanlang (1-7):{Colors.END} ").strip()
        
        if choice == '7':
            self.export_navoiy_format()
            return
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = os.path.splitext(os.path.basename(self.current_pdf_path))[0]
        
        formats = {
            '1': ('.txt', ExportManager.to_txt),
            '2': ('.json', ExportManager.to_json),
            '3': ('.csv', ExportManager.to_csv),
            '4': ('.html', ExportManager.to_html),
            '5': ('.docx', ExportManager.to_word),
            '6': ('.xlsx', ExportManager.to_excel)
        }
        
        if choice in formats:
            ext, func = formats[choice]
            
            if choice in ['5'] and not DOCX_AVAILABLE:
                print(f"{Colors.RED}❌ DOCX qo'llab-quvvatlanmaydi. pip install python-docx{Colors.END}")
                return
            
            if choice in ['6'] and not EXCEL_AVAILABLE:
                print(f"{Colors.RED}❌ XLSX qo'llab-quvvatlanmaydi. pip install openpyxl{Colors.END}")
                return
            
            output_path = f"{base_name}_export_{timestamp}{ext}"
            
            spinner = Spinner(f"{ext} formatida saqlanmoqda...")
            spinner.start()
            
            if func(self.current_result, output_path):
                spinner.stop(f"✅ Saqlandi: {output_path}")
            else:
                spinner.stop(f"❌ Saqlashda xatolik!")
        else:
            print(f"{Colors.RED}❌ Noto'g'ri tanlov!{Colors.END}")
    
    def ocr_extract(self):
        if not self.current_pdf_path:
            print(f"{Colors.YELLOW}⚠️ Avval PDF fayl tanlang!{Colors.END}")
            return
        
        ocr = OCRExtractor()
        if not ocr.available:
            return
        
        print(f"\n{Colors.CYAN}🖼️ OCR matn olish boshlanmoqda...{Colors.END}")
        print(f"{Colors.GRAY}Bu jarayon PDF hajmiga qarab bir necha daqiqa olishi mumkin.{Colors.END}")
        
        lang = input(f"{Colors.CYAN}Tillar (uzb+rus+eng):{Colors.END} ").strip() or 'uzb+rus+eng'
        
        text = ocr.extract(self.current_pdf_path, lang)
        
        if text:
            save = input(f"\n{Colors.CYAN}💾 Saqlash kerakmi? (ha/yo'q):{Colors.END} ").lower()
            if save in ['ha', 'h', 'yes', 'y']:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_path = f"{os.path.splitext(os.path.basename(self.current_pdf_path))[0]}_ocr_{timestamp}.txt"
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                print(f"{Colors.GREEN}✅ Saqlandi: {output_path}{Colors.END}")
        else:
            print(f"{Colors.RED}❌ OCR bajarilmadi!{Colors.END}")
    
    def security_menu(self):
        print(f"\n{Colors.BOLD}🔐 XAVFSIZLIK AMALLARI:{Colors.END}")
        print(f"{Colors.GREEN}1.{Colors.END} Shifrlash kaliti yaratish")
        print(f"{Colors.GREEN}2.{Colors.END} Faylni shifrlash")
        print(f"{Colors.GREEN}3.{Colors.END} Faylni deshifrlash")
        print(f"{Colors.GREEN}4.{Colors.END} Fayl xeshini hisoblash")
        
        if not CRYPTO_AVAILABLE:
            print(f"\n{Colors.YELLOW}⚠️ cryptography kutubxonasi o'rnatilmagan!{Colors.END}")
            print(f"{Colors.GRAY}   pip install cryptography{Colors.END}")
            return
        
        choice = input(f"\n{Colors.CYAN}👉 Tanlang (1-4):{Colors.END} ").strip()
        
        if choice == '1':
            key = SecurityManager.generate_key()
            print(f"\n{Colors.GREEN}✅ Shifrlash kaliti:{Colors.END}")
            print(f"{Colors.YELLOW}{key.decode()}{Colors.END}")
            print(f"{Colors.RED}⚠️ Bu kalitni saqlab qo'ying!{Colors.END}")
            
            save = input(f"\n{Colors.CYAN}Kalitni faylga saqlash? (ha/yo'q):{Colors.END} ").lower()
            if save in ['ha', 'h', 'yes', 'y']:
                with open('encryption_key.key', 'wb') as f:
                    f.write(key)
                print(f"{Colors.GREEN}✅ Kalit 'encryption_key.key' faylida saqlandi{Colors.END}")
        
        elif choice == '2':
            file_path = input(f"{Colors.CYAN}Shifrlanadigan fayl:{Colors.END} ").strip()
            if not os.path.exists(file_path):
                print(f"{Colors.RED}❌ Fayl topilmadi!{Colors.END}")
                return
            
            key_input = input(f"{Colors.CYAN}Shifrlash kaliti (yoki fayl nomi):{Colors.END} ").strip()
            
            if os.path.exists(key_input):
                with open(key_input, 'rb') as f:
                    key = f.read()
            else:
                key = key_input.encode()
            
            output = file_path + '.encrypted'
            if SecurityManager.encrypt_file(file_path, output, key):
                print(f"{Colors.GREEN}✅ Shifrlandi: {output}{Colors.END}")
            else:
                print(f"{Colors.RED}❌ Shifrlashda xatolik!{Colors.END}")
        
        elif choice == '3':
            file_path = input(f"{Colors.CYAN}Shifrlangan fayl:{Colors.END} ").strip()
            if not os.path.exists(file_path):
                print(f"{Colors.RED}❌ Fayl topilmadi!{Colors.END}")
                return
            
            key_input = input(f"{Colors.CYAN}Shifrlash kaliti:{Colors.END} ").strip()
            
            if os.path.exists(key_input):
                with open(key_input, 'rb') as f:
                    key = f.read()
            else:
                key = key_input.encode()
            
            output = file_path.replace('.encrypted', '.decrypted')
            if SecurityManager.decrypt_file(file_path, output, key):
                print(f"{Colors.GREEN}✅ Deshifrlandi: {output}{Colors.END}")
            else:
                print(f"{Colors.RED}❌ Deshifrlashda xatolik!{Colors.END}")
        
        elif choice == '4':
            file_path = input(f"{Colors.CYAN}Fayl manzili:{Colors.END} ").strip()
            if os.path.exists(file_path):
                alg = input(f"{Colors.CYAN}Algoritm (md5/sha1/sha256) [sha256]:{Colors.END} ").strip() or 'sha256'
                hash_val = SecurityManager.get_file_hash(file_path, alg)
                print(f"\n{Colors.GREEN}{alg.upper()}:{Colors.END} {hash_val}")
            else:
                print(f"{Colors.RED}❌ Fayl topilmadi!{Colors.END}")
    
    def settings_menu(self):
        print(f"\n{Colors.BOLD}⚙️ SOZLAMALAR:{Colors.END}")
        print(f"{Colors.GREEN}1.{Colors.END} Workerlar sonini o'zgartirish (hozir: {self.max_workers})")
        print(f"{Colors.GREEN}2.{Colors.END} Kutubxonalar holatini tekshirish")
        
        choice = input(f"\n{Colors.CYAN}👉 Tanlang:{Colors.END} ").strip()
        
        if choice == '1':
            try:
                cpu_count = os.cpu_count() or 4
                new_workers = int(input(f"{Colors.CYAN}Workerlar soni (1-{cpu_count*2}):{Colors.END} "))
                if 1 <= new_workers <= cpu_count * 2:
                    self.max_workers = new_workers
                    print(f"{Colors.GREEN}✅ Workerlar soni: {self.max_workers}{Colors.END}")
                else:
                    print(f"{Colors.RED}❌ Noto'g'ri qiymat!{Colors.END}")
            except ValueError:
                print(f"{Colors.RED}❌ Noto'g'ri format!{Colors.END}")
        
        elif choice == '2':
            print(f"\n{Colors.BOLD}📚 KUTUBXONALAR HOLATI:{Colors.END}")
            print(f"   pdfplumber: {Colors.GREEN}✓{Colors.END}")
            print(f"   cryptography: {Colors.GREEN if CRYPTO_AVAILABLE else Colors.RED}{'✓' if CRYPTO_AVAILABLE else '✗'}{Colors.END}")
            print(f"   python-docx: {Colors.GREEN if DOCX_AVAILABLE else Colors.RED}{'✓' if DOCX_AVAILABLE else '✗'}{Colors.END}")
            print(f"   openpyxl: {Colors.GREEN if EXCEL_AVAILABLE else Colors.RED}{'✓' if EXCEL_AVAILABLE else '✗'}{Colors.END}")
            print(f"   googletrans: {Colors.GREEN if TRANSLATOR_AVAILABLE else Colors.RED}{'✓' if TRANSLATOR_AVAILABLE else '✗'}{Colors.END}")
            print(f"   pytesseract+pdf2image: {Colors.GREEN if OCR_AVAILABLE else Colors.RED}{'✓' if OCR_AVAILABLE else '✗'}{Colors.END}")


# ==================== DASTURNI ISHGA TUSHIRISH ====================
def main():
    """Asosiy funksiya"""
    if sys.version_info < (3, 8):
        print(f"{Colors.RED}❌ Python 3.8 yoki undan yuqori versiya kerak!{Colors.END}")
        sys.exit(1)
    
    try:
        cli = PDFExtractorCLI()
        cli.run()
    except KeyboardInterrupt:
        print(f"\n\n{Colors.YELLOW}⚠️ Dastur to'xtatildi.{Colors.END}")
    except Exception as e:
        print(f"\n{Colors.RED}❌ Kutilmagan xatolik: {e}{Colors.END}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()